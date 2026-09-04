"""Compare the revision's models against the numbers printed in the manuscript.

Usage: python scripts/build_genomewide_vs_region_tables.py [supplementary.docx]
Writes Results/genomewide_vs_region_models.csv (never overwrites originals).
"""
import os
import sys

import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import analysis_cells

try:
    import docx
except ImportError:
    sys.exit("python-docx is required: pip install --user python-docx")

import manuscript_supp

SUPP = manuscript_supp.find(sys.argv[1] if len(sys.argv) > 1 else None)
if SUPP is None:
    sys.exit("supplementary DOCX not found; pass its path as the first argument")

ANALYSIS, OUTDIR = "analysis", "Results"
RANDS, MISSING = [0, 42, 1024], ["5%", "15%", "25%"]
SEG = "seg1024_overlap128"
# manuscript bin labels, in the column order the supplementary tables use
MS_BINS = ["<=0.5%", "0.5%~1%", "1%~10%", "10%~20%", "20%~30%", "30%~40%",
           "40%~50%", "Overall"]
# the revision's bottom bin starts at 0.1%; every other label matches
BIN_ALIAS = {"0.1%~0.5%": "<=0.5%"}

# anchor -> (dataset, chromosome, runId, epoch, {missing: docx table index}).
# The epoch is the checkpoint the paper reports; None means untagged cells.
ANCHORS = {
    "1KGP_chr22": ("1KGP", 22, "1KGP_genomewide", 70, {"5%": 0, "15%": 1, "25%": 2}),
    "SGDP_chr22": ("SGDP", 22, "SGDP_genomewide", "best", {"5%": 6, "15%": 7, "25%": 8}),
    "SGDP_chr19": ("SGDP", 19, "SGDP_genomewide", "best", {"5%": 12, "15%": 13, "25%": 14}),
}

doc = docx.Document(SUPP)


def manuscript_rows(table_index):
    """{(metric, model): [7 bins + overall]} from one supplementary table."""
    t = doc.tables[table_index]
    out = {}
    for row in t.rows:
        cells = [c.text.strip() for c in row.cells]
        if cells[0] in ("Acc", "R2") and cells[1] in ("BiU-Net", "Beagle", "SCDA"):
            try:
                out[(cells[0], cells[1])] = [float(x) for x in cells[2:10]]
            except ValueError:
                continue
    return out


# Both denominators are reported: the markers Beagle retains, and the full
# target axis. Marker counts are carried alongside the values.
DENOMINATORS = (("BAT", "Beagle intersection"), ("BAT_FA", "full target axis"))


def revision_rows(dataset, chrom, run_id, epoch, miss, flags="BAT"):
    """Seed-averaged Bin_R2/Bin_Acc per bin, plus the Overall row."""
    used, frames = [], []
    for rs in RANDS:
        p = analysis_cells.resolve(ANALYSIS, run_id, "phased", "biunet", flags,
                                   rs, dataset, chrom, "ALL", SEG, miss,
                                   epoch=epoch)
        if p is None:
            continue
        used.append(p)
        frames.append(pd.read_csv(p))
    if not frames:
        return None, []
    df = pd.concat(frames)
    df["MAF_bin"] = df["MAF_bin"].astype(str).str.strip().replace(BIN_ALIAS)
    agg = df.groupby("MAF_bin").agg(R2=("Bin_R2", "mean"), Acc=("Bin_Acc", "mean"),
                                    N=("Num_SNPs", "mean"))
    vals = {b: (agg.loc[b, "R2"], agg.loc[b, "Acc"], agg.loc[b, "N"]) for b in agg.index}
    vals["Overall"] = (df["Overall R2"].mean(), df["Overall Acc"].mean(),
                       df[df.MAF_bin == "Overall"]["Num_SNPs"].mean())
    return vals, used


def manuscript_nsnps(table_index):
    """Marker count behind each published bin, from the table's #SNPs row."""
    for row in doc.tables[table_index].rows:
        c = [x.text.strip() for x in row.cells]
        if c[0] == "#SNPs" and c[1] == "#SNPs":
            try:
                return [int(x.replace(",", "")) for x in c[2:10]]
            except ValueError:
                return None
    return None


rows, provenance = [], []
for label, (dataset, chrom, run_id, epoch, tmap) in ANCHORS.items():
    for flags, denom in DENOMINATORS:
        for miss, tidx in tmap.items():
            ms = manuscript_rows(tidx)
            ns = manuscript_nsnps(tidx)
            rev, used = revision_rows(dataset, chrom, run_id, epoch, miss, flags=flags)
            provenance.extend(used)
            if rev is None:
                if flags == "BAT":
                    print(f"WARNING: no revision cells for {label} {miss} ({run_id} epoch={epoch})")
                continue
            for i, b in enumerate(MS_BINS):
                if b not in rev:
                    continue
                r2_new, acc_new, n_new = rev[b]
                rows.append({
                    "anchor": label, "denominator": denom, "missing": miss, "MAF_bin": b,
                    "R2_manuscript": ms[("R2", "BiU-Net")][i],
                    "R2_revision": r2_new,
                    "R2_delta": r2_new - ms[("R2", "BiU-Net")][i],
                    "Acc_manuscript": ms[("Acc", "BiU-Net")][i],
                    "Acc_revision": acc_new,
                    "Acc_delta": acc_new - ms[("Acc", "BiU-Net")][i],
                    "R2_beagle_manuscript": ms[("R2", "Beagle")][i],
                    "Num_SNPs_revision": n_new,
                    "Num_SNPs_manuscript": ns[i] if ns else None,
                    "same_variant_set": (None if not ns
                                         else int(round(n_new)) == int(ns[i])),
                    "revision_run": run_id,
                    "model_scope": "genome-wide single model",
                })

if not rows:
    sys.exit("No anchors could be built — no revision cells found.")

os.makedirs(OUTDIR, exist_ok=True)
out = pd.DataFrame(rows)
path = os.path.join(OUTDIR, "genomewide_vs_region_models.csv")
out.to_csv(path, index=False)
print(f"wrote {path} ({len(out)} rows)")
print(f"Revision cells: {analysis_cells.describe(provenance)}")
print(f"Manuscript values parsed from: {os.path.basename(SUPP)}")
for label in ANCHORS:
    for _, denom in DENOMINATORS:
        sub = out[(out.anchor == label) & (out.MAF_bin == "Overall")
                  & (out.denominator == denom)]
        if not sub.empty:
            deltas = ", ".join(f"{m} {d:+.4f}" for m, d in
                               zip(sub.missing, sub.R2_delta))
            print(f"  {label} [{denom}] Overall R2 vs published: {deltas}")
